"""Generate PDF / Word / CSV reports from a brand bundle.

PDF / Word now includes:
  • Branded cover with colored band, KPI tiles, completion donut, readiness radar
  • Executive summary narrative
  • Charts: fields per section, mandatory %, integration %, source mix
  • Conflicts to Resolve (auto-detected, multi-axis)
  • Insights & Recommendations (rules-based, color-coded)
  • Per-section detail with new Integration / Source columns
"""
from io import BytesIO
import pandas as pd
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import cm
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Image,
)
from docx import Document
from docx.shared import Pt, RGBColor, Inches

import analysis as A

SECTIONS = [
    ("people",                "A. People & Stakeholders"),
    ("partner_360",           "B. Partner 360 — Fields"),
    ("customer_360",          "C. Customer 360 — Fields"),
    ("sales_opp_details",     "D1. Sales — Opportunity Details"),
    ("sales_contact_details", "D2. Sales — Contact Details"),
    ("sales_deal_details",    "D3. Sales — Deal Details"),
    ("approvals",             "E. Approval Stages & Workflow"),
    ("dashboards",            "F. Dashboards & Reports Expected"),
    ("best_practices",        "G. Best-Practice Inputs"),
    ("open_notes",            "H. Open Notes / Pain Points / Asks"),
]
FIELD_BUILDER_SECTIONS = A.FIELD_BUILDER_SECTIONS

RED   = colors.HexColor("#C8102E")
DARK  = colors.HexColor("#1A1A1A")
GREY  = colors.HexColor("#666666")
LIGHT = colors.HexColor("#F5F5F5")
ACCENT_BLUE = colors.HexColor("#0F4C81")


def _safe(v) -> str:
    if v is None: return ""
    if isinstance(v, list): return ", ".join(str(x) for x in v if x)
    if isinstance(v, bool): return "Yes" if v else "No"
    return str(v)


# =====================================================================
# CSV
# =====================================================================

def build_csv(bundle: dict) -> bytes:
    rows = []
    brand = bundle["brand"]
    for c in bundle["contributors"]:
        contrib_label = f"{c['name']} ({c['role']})"
        sections = bundle["responses_by_contributor"].get(c["id"], {})
        for sec_key, sec_title in SECTIONS:
            payload = sections.get(sec_key, {})
            if not payload: continue
            if sec_key in FIELD_BUILDER_SECTIONS:
                for f in payload.get("fields", []) or []:
                    rows.append({
                        "Brand": brand, "Contributor": contrib_label, "Section": sec_title,
                        "Field": f.get("field", ""), "Type": f.get("type", ""),
                        "Options": f.get("options", ""), "Mandatory": _safe(f.get("mandatory")),
                        "Integration Needed": _safe(f.get("integration_needed")),
                        "Data Source": f.get("data_capture_source", ""),
                        "Conditional Rule": f.get("conditional_rule", ""), "Value": "",
                    })
            elif sec_key == "approvals":
                for s in payload.get("stages", []) or []:
                    rows.append({
                        "Brand": brand, "Contributor": contrib_label, "Section": sec_title,
                        "Field": s.get("stage", ""), "Type": "Approval Stage",
                        "Options": s.get("approver", ""), "Mandatory": s.get("trigger", ""),
                        "Integration Needed": "", "Data Source": "",
                        "Conditional Rule": f"SLA {s.get('sla_hours','')}h, can_revert={s.get('can_revert','')}",
                        "Value": "",
                    })
            elif sec_key == "dashboards":
                for d in payload.get("dashboards", []) or []:
                    rows.append({
                        "Brand": brand, "Contributor": contrib_label, "Section": sec_title,
                        "Field": d.get("dashboard", ""), "Type": "Dashboard",
                        "Options": d.get("audience", ""), "Mandatory": d.get("frequency", ""),
                        "Integration Needed": "", "Data Source": "",
                        "Conditional Rule": "", "Value": "",
                    })
            else:
                for k, v in payload.items():
                    rows.append({
                        "Brand": brand, "Contributor": contrib_label, "Section": sec_title,
                        "Field": k, "Type": "", "Options": "", "Mandatory": "",
                        "Integration Needed": "", "Data Source": "",
                        "Conditional Rule": "", "Value": _safe(v),
                    })
    cols = ["Brand","Contributor","Section","Field","Type","Options","Mandatory","Integration Needed","Data Source","Conditional Rule","Value"]
    df = pd.DataFrame(rows) if rows else pd.DataFrame(columns=cols)
    return df.to_csv(index=False).encode("utf-8")


def build_cross_brand_csv(bundles: list[dict]) -> bytes:
    frames = []
    for b in bundles:
        frames.append(pd.read_csv(BytesIO(build_csv(b))))
    if not frames:
        return b""
    return pd.concat(frames, ignore_index=True).to_csv(index=False).encode("utf-8")


# =====================================================================
# PDF — styles + components
# =====================================================================

def _styles():
    s = getSampleStyleSheet()
    return {
        "h_cover":  ParagraphStyle("h_cover", parent=s["Heading1"], fontName="Helvetica-Bold", fontSize=30, textColor=colors.white, spaceAfter=8, leading=34),
        "sub_cover":ParagraphStyle("sub_cover", parent=s["BodyText"], fontSize=12, textColor=colors.white),
        "kicker":   ParagraphStyle("kicker", parent=s["BodyText"], fontSize=10, textColor=colors.white, spaceAfter=4),
        "h1":       ParagraphStyle("h1", parent=s["Heading1"], fontName="Helvetica-Bold", fontSize=18, textColor=RED, spaceAfter=10, spaceBefore=4),
        "h1_white": ParagraphStyle("h1_white", parent=s["Heading1"], fontName="Helvetica-Bold", fontSize=18, textColor=colors.white, spaceAfter=10),
        "h2":       ParagraphStyle("h2", parent=s["Heading2"], fontSize=13, textColor=DARK, spaceAfter=6),
        "h3":       ParagraphStyle("h3", parent=s["Heading3"], fontSize=11, textColor=GREY, spaceAfter=4),
        "body":     ParagraphStyle("body", parent=s["BodyText"], fontSize=10, leading=14, textColor=DARK),
        "small":    ParagraphStyle("small", parent=s["BodyText"], fontSize=8, textColor=GREY),
        "tile_v":   ParagraphStyle("tile_v", parent=s["BodyText"], fontName="Helvetica-Bold", fontSize=20, textColor=RED, alignment=1),
        "tile_l":   ParagraphStyle("tile_l", parent=s["BodyText"], fontSize=8, textColor=GREY, alignment=1),
        "callout":  ParagraphStyle("callout", parent=s["BodyText"], fontSize=10, textColor=DARK, leftIndent=8, borderPadding=4),
    }


def _cover_band(brand: str, generated_at: str, st):
    """Top color band on the cover page (full-width red bar with title)."""
    band = Table(
        [[Paragraph("REDINGTON  ·  Cloud & AI Practice", st["kicker"])],
         [Paragraph("Zoho CRM Discovery Pack", st["h_cover"])],
         [Paragraph(f"Brand: <b>{brand}</b>", st["sub_cover"])],
         [Paragraph(f"Generated: {generated_at}", st["sub_cover"])]],
        colWidths=[17*cm],
    )
    band.setStyle(TableStyle([
        ("BACKGROUND", (0,0), (-1,-1), RED),
        ("LEFTPADDING", (0,0), (-1,-1), 16),
        ("RIGHTPADDING", (0,0), (-1,-1), 16),
        ("TOPPADDING",  (0,0), (-1,-1), 6),
        ("BOTTOMPADDING", (0,0), (-1,-1), 4),
        ("VALIGN", (0,0), (-1,-1), "TOP"),
    ]))
    return band


def _kpi_tiles(metrics: dict, st) -> Table:
    items = [
        (str(metrics["contributors"]),     "Contributors"),
        (str(metrics["total_fields"]),     "Fields proposed"),
        (f"{metrics['mandatory_pct']}%",   "Mandatory %"),
        (f"{metrics['integration_pct']}%", "Need integration"),
        (str(metrics["approval_stages"]),  "Approval stages"),
        (str(metrics["dashboards"]),       "Dashboards"),
    ]
    cells = [
        [Paragraph(v, st["tile_v"]) for v, _ in items],
        [Paragraph(l, st["tile_l"]) for _, l in items],
    ]
    t = Table(cells, colWidths=[2.83*cm]*6)
    t.setStyle(TableStyle([
        ("BACKGROUND", (0,0), (-1,-1), LIGHT),
        ("BOX",        (0,0), (-1,-1), 0.5, GREY),
        ("INNERGRID",  (0,0), (-1,-1), 0.25, colors.HexColor("#DDD")),
        ("VALIGN",     (0,0), (-1,-1), "MIDDLE"),
        ("ALIGN",      (0,0), (-1,-1), "CENTER"),
        ("TOPPADDING", (0,0), (-1,-1), 8),
        ("BOTTOMPADDING", (0,0), (-1,-1), 8),
    ]))
    return t


def _section_header(title: str, st):
    """Colored sidebar block + section heading."""
    bar = Table([[" "]], colWidths=[0.18*cm], rowHeights=[0.7*cm])
    bar.setStyle(TableStyle([("BACKGROUND", (0,0), (-1,-1), RED)]))
    head = Paragraph(title, st["h1"])
    t = Table([[bar, head]], colWidths=[0.4*cm, 16.6*cm])
    t.setStyle(TableStyle([("VALIGN", (0,0), (-1,-1), "MIDDLE"), ("LEFTPADDING",(0,0),(-1,-1),0)]))
    return t


def _recommendations_table(recs: list[dict]) -> Table | None:
    if not recs: return None
    rows = [["Priority", "Topic", "Recommendation"]]
    for r in recs: rows.append([r["priority"], r["topic"], r["text"]])
    t = Table(rows, colWidths=[2*cm, 3.5*cm, 11.5*cm], repeatRows=1)
    style = [
        ("BACKGROUND", (0,0), (-1,0), DARK),
        ("TEXTCOLOR",  (0,0), (-1,0), colors.whitesmoke),
        ("FONTNAME",   (0,0), (-1,0), "Helvetica-Bold"),
        ("FONTSIZE",   (0,0), (-1,-1), 8),
        ("GRID",       (0,0), (-1,-1), 0.25, GREY),
        ("VALIGN",     (0,0), (-1,-1), "TOP"),
        ("LEFTPADDING",(0,0), (-1,-1), 5),
        ("RIGHTPADDING",(0,0), (-1,-1), 5),
        ("TOPPADDING", (0,0), (-1,-1), 4),
        ("BOTTOMPADDING", (0,0), (-1,-1), 4),
    ]
    bg_map = {"High": colors.HexColor("#FCE3E6"), "Medium": colors.HexColor("#FFF4E0"), "Low": colors.HexColor("#E8F4FD"), "Info": colors.HexColor("#EAF7EE")}
    fg_map = {"High": colors.HexColor("#A50D26"), "Medium": colors.HexColor("#8A5A00"), "Low": colors.HexColor("#0F4C81"), "Info": colors.HexColor("#2E7D32")}
    for i, r in enumerate(recs, start=1):
        style.append(("BACKGROUND", (0,i), (0,i), bg_map.get(r["priority"], colors.white)))
        style.append(("TEXTCOLOR",  (0,i), (0,i), fg_map.get(r["priority"], DARK)))
        style.append(("FONTNAME",   (0,i), (0,i), "Helvetica-Bold"))
    t.setStyle(TableStyle(style))
    return t


def _conflicts_table(conflicts: list[dict]) -> Table | None:
    if not conflicts: return None
    rows = [["Section", "Type", "Item", "Detail"]]
    for c in conflicts:
        rows.append([c["section"], c["type"], c["item"], c["detail"]])
    t = Table(rows, colWidths=[3.5*cm, 3.2*cm, 3.5*cm, 6.8*cm], repeatRows=1)
    t.setStyle(TableStyle([
        ("BACKGROUND", (0,0), (-1,0), RED),
        ("TEXTCOLOR",  (0,0), (-1,0), colors.whitesmoke),
        ("FONTNAME",   (0,0), (-1,0), "Helvetica-Bold"),
        ("FONTSIZE",   (0,0), (-1,-1), 8),
        ("GRID",       (0,0), (-1,-1), 0.25, GREY),
        ("VALIGN",     (0,0), (-1,-1), "TOP"),
    ]))
    return t


def _page_decoration(canvas, doc):
    canvas.saveState()
    canvas.setFillColor(RED)
    canvas.rect(0, 0, A4[0], 0.18*cm, fill=1, stroke=0)
    canvas.setFillColor(GREY)
    canvas.setFont("Helvetica", 7.5)
    canvas.drawString(2*cm, 0.6*cm, "Redington · Zoho CRM Discovery")
    canvas.drawRightString(A4[0] - 2*cm, 0.6*cm, f"Page {doc.page}")
    canvas.restoreState()


# =====================================================================
# PDF
# =====================================================================

def build_pdf(bundle: dict) -> bytes:
    metrics    = A.compute_metrics(bundle)
    conflicts  = A.detect_conflicts(bundle)
    recs       = A.generate_recommendations(bundle, metrics, conflicts)

    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4, leftMargin=2*cm, rightMargin=2*cm, topMargin=1.5*cm, bottomMargin=1.5*cm,
        title=f"Redington Zoho CRM Discovery — {bundle['brand']}",
        author="Redington Cloud & AI Practice",
    )
    st = _styles()
    story = []

    # ---------- COVER ----------
    story.append(_cover_band(bundle["brand"], bundle["generated_at"], st))
    story.append(Spacer(1, 0.6*cm))

    # KPI strip
    story.append(_kpi_tiles(metrics, st))
    story.append(Spacer(1, 0.5*cm))

    # Donut + Radar side-by-side
    try:
        donut  = Image(BytesIO(A.chart_completion_donut_png(metrics)),  width=5*cm, height=5*cm)
        radar  = Image(BytesIO(A.chart_readiness_radar_png(metrics)),    width=8*cm, height=8*cm)
        side   = Table([[donut, radar]], colWidths=[5.5*cm, 11.5*cm])
        side.setStyle(TableStyle([("VALIGN", (0,0), (-1,-1), "MIDDLE")]))
        story.append(side)
    except Exception:
        pass
    story.append(Spacer(1, 0.4*cm))

    # Contributors mini-table
    story.append(Paragraph("Contributors", st["h2"]))
    contrib_rows = [["Name", "Role", "Email", "Last update"]]
    for c in bundle["contributors"]:
        contrib_rows.append([c["name"], c["role"], c["email"], c["submitted_at"][:19].replace("T", " ")])
    if len(contrib_rows) == 1:
        contrib_rows.append(["—", "—", "—", "—"])
    contrib_t = Table(contrib_rows, colWidths=[3.6*cm, 2.4*cm, 5.8*cm, 5.2*cm], repeatRows=1)
    contrib_t.setStyle(TableStyle([
        ("BACKGROUND", (0,0), (-1,0), DARK),
        ("TEXTCOLOR",  (0,0), (-1,0), colors.whitesmoke),
        ("FONTNAME",   (0,0), (-1,0), "Helvetica-Bold"),
        ("FONTSIZE",   (0,0), (-1,-1), 8),
        ("GRID",       (0,0), (-1,-1), 0.25, GREY),
        ("VALIGN",     (0,0), (-1,-1), "TOP"),
    ]))
    story.append(contrib_t)

    # ---------- EXEC SUMMARY ----------
    story.append(PageBreak())
    story.append(_section_header("Executive Summary", st))
    story.append(Paragraph(
        f"This pack captures the Zoho CRM discovery for the <b>{bundle['brand']}</b> practice. "
        f"<b>{metrics['contributors']}</b> contributor(s) proposed <b>{metrics['total_fields']}</b> fields across "
        f"{len(FIELD_BUILDER_SECTIONS)} field-builder sections, with <b>{metrics['mandatory_pct']}%</b> marked mandatory and "
        f"<b>{metrics['integration_pct']}%</b> flagged as needing integration. "
        f"Data is sourced from <b>{metrics['unique_sources']}</b> different system(s). "
        f"<b>{metrics['approval_stages']}</b> approval stage(s) and <b>{metrics['dashboards']}</b> dashboard(s) are requested. "
        f"Section-completion is currently <b>{metrics['completion_pct']}%</b>; overall brand-readiness scores "
        f"<b>{A.overall_readiness_pct(metrics):.0f}/100</b>.",
        st["body"]))
    story.append(Spacer(1, 0.4*cm))

    try:
        story.append(Image(BytesIO(A.chart_fields_per_section_png(bundle)),       width=17*cm, height=7.6*cm))
        story.append(Spacer(1, 0.25*cm))
        story.append(Image(BytesIO(A.chart_mandatory_pct_png(bundle)),             width=17*cm, height=7.6*cm))
        story.append(PageBreak())
        story.append(_section_header("Integration & Data Sources", st))
        story.append(Paragraph(
            "How much of the data model needs integration, and where does the data live today? "
            "Use this to plan Phase-2 integrations and define source-of-truth per entity.", st["small"]))
        story.append(Spacer(1, 0.25*cm))
        story.append(Image(BytesIO(A.chart_integration_per_section_png(bundle)),   width=17*cm, height=7.6*cm))
        story.append(Spacer(1, 0.3*cm))
        story.append(Image(BytesIO(A.chart_source_breakdown_png(bundle)),          width=14*cm, height=8.5*cm))
    except Exception:
        pass

    # ---------- CONFLICTS ----------
    if conflicts:
        story.append(PageBreak())
        story.append(_section_header("Conflicts to Resolve", st))
        story.append(Paragraph(
            "Items where contributors disagree (mandatory flag, data source, integration need, or SLA). "
            "Reconcile before Zoho handover — these are decision points, not data quality issues.", st["small"]))
        story.append(Spacer(1, 0.2*cm))
        story.append(_conflicts_table(conflicts))

    # ---------- RECOMMENDATIONS ----------
    story.append(PageBreak())
    story.append(_section_header("Insights & Recommendations", st))
    story.append(Paragraph(
        "Auto-generated based on what the brand team has filled in. "
        "Color-coded by priority: <b>High</b> = act before handover, <b>Medium</b> = address in working session, "
        "<b>Low</b> = consider/validate, <b>Info</b> = healthy state.", st["small"]))
    story.append(Spacer(1, 0.2*cm))
    rec_t = _recommendations_table(recs)
    if rec_t: story.append(rec_t)

    # ---------- PER-SECTION DETAIL ----------
    for sec_key, sec_title in SECTIONS:
        story.append(PageBreak())
        story.append(_section_header(sec_title, st))
        any_data = False
        for c in bundle["contributors"]:
            payload = bundle["responses_by_contributor"].get(c["id"], {}).get(sec_key)
            if not payload: continue
            any_data = True
            story.append(Paragraph(f"From {c['name']} ({c['role']})", st["h3"]))
            _render_section_pdf(story, sec_key, payload, st)
            story.append(Spacer(1, 0.3*cm))
        if not any_data:
            story.append(Paragraph("<i>No input captured for this section.</i>", st["body"]))

    doc.build(story, onFirstPage=_page_decoration, onLaterPages=_page_decoration)
    return buf.getvalue()


def _render_section_pdf(story, sec_key, payload, st):
    if sec_key in FIELD_BUILDER_SECTIONS:
        rows = payload.get("fields", []) or []
        if rows:
            tab = [["Field", "Type", "Mand", "Integ", "Source", "Options", "Conditional Rule"]]
            for f in rows:
                tab.append([
                    _safe(f.get("field")), _safe(f.get("type")),
                    "Y" if f.get("mandatory") else "N",
                    "Y" if f.get("integration_needed") else "N",
                    _safe(f.get("data_capture_source")),
                    _safe(f.get("options")),
                    _safe(f.get("conditional_rule")),
                ])
            t = Table(tab, colWidths=[3.4*cm, 1.5*cm, 0.9*cm, 0.9*cm, 2.0*cm, 3.6*cm, 4.7*cm], repeatRows=1)
            t.setStyle(TableStyle([
                ("BACKGROUND", (0,0), (-1,0), DARK),
                ("TEXTCOLOR",  (0,0), (-1,0), colors.whitesmoke),
                ("FONTNAME",   (0,0), (-1,0), "Helvetica-Bold"),
                ("FONTSIZE",   (0,0), (-1,-1), 6.8),
                ("GRID",       (0,0), (-1,-1), 0.25, GREY),
                ("VALIGN",     (0,0), (-1,-1), "TOP"),
                ("ALIGN",      (2,1), (3,-1), "CENTER"),
                ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.white, LIGHT]),
            ]))
            story.append(t)
        if payload.get("notes"):
            story.append(Spacer(1, 0.15*cm))
            story.append(Paragraph(f"<b>Notes:</b> {_safe(payload['notes'])}", st["small"]))

    elif sec_key == "approvals":
        rows = payload.get("stages", []) or []
        if rows:
            tab = [["Stage", "Approver", "Trigger", "SLA (h)", "Revert?"]]
            for s in rows:
                tab.append([_safe(s.get("stage")), _safe(s.get("approver")), _safe(s.get("trigger")),
                            _safe(s.get("sla_hours")), "Yes" if s.get("can_revert") else "No"])
            t = Table(tab, colWidths=[3.8*cm, 3.2*cm, 5.5*cm, 1.7*cm, 2.3*cm], repeatRows=1)
            t.setStyle(TableStyle([
                ("BACKGROUND", (0,0), (-1,0), DARK),
                ("TEXTCOLOR",  (0,0), (-1,0), colors.whitesmoke),
                ("FONTNAME",   (0,0), (-1,0), "Helvetica-Bold"),
                ("FONTSIZE",   (0,0), (-1,-1), 7),
                ("GRID",       (0,0), (-1,-1), 0.25, GREY),
                ("VALIGN",     (0,0), (-1,-1), "TOP"),
                ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.white, LIGHT]),
            ]))
            story.append(t)
        for k in ("escalation", "notes"):
            if payload.get(k):
                story.append(Paragraph(f"<b>{k.title()}:</b> {_safe(payload[k])}", st["small"]))

    elif sec_key == "dashboards":
        rows = payload.get("dashboards", []) or []
        if rows:
            tab = [["Dashboard", "Audience", "Frequency"]]
            for d in rows:
                tab.append([_safe(d.get("dashboard")), _safe(d.get("audience")), _safe(d.get("frequency"))])
            t = Table(tab, colWidths=[7.5*cm, 6*cm, 3*cm], repeatRows=1)
            t.setStyle(TableStyle([
                ("BACKGROUND", (0,0), (-1,0), DARK),
                ("TEXTCOLOR",  (0,0), (-1,0), colors.whitesmoke),
                ("FONTNAME",   (0,0), (-1,0), "Helvetica-Bold"),
                ("FONTSIZE",   (0,0), (-1,-1), 7),
                ("GRID",       (0,0), (-1,-1), 0.25, GREY),
                ("VALIGN",     (0,0), (-1,-1), "TOP"),
                ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.white, LIGHT]),
            ]))
            story.append(t)
        if payload.get("notes"):
            story.append(Paragraph(f"<b>Notes:</b> {_safe(payload['notes'])}", st["small"]))

    else:
        for k, v in payload.items():
            if v:
                story.append(Paragraph(f"<b>{k.replace('_',' ').title()}:</b> {_safe(v)}", st["body"]))


# =====================================================================
# Word
# =====================================================================

def _docx_set_heading_color(par, hex_color: str):
    for run in par.runs:
        run.font.color.rgb = RGBColor.from_string(hex_color.lstrip("#"))


def build_docx(bundle: dict) -> bytes:
    metrics   = A.compute_metrics(bundle)
    conflicts = A.detect_conflicts(bundle)
    recs      = A.generate_recommendations(bundle, metrics, conflicts)

    doc = Document()

    # Cover
    p = doc.add_paragraph()
    r = p.add_run("REDINGTON  ·  Cloud & AI Practice")
    r.font.size = Pt(11); r.font.color.rgb = RGBColor(0xC8, 0x10, 0x2E); r.bold = True

    h = doc.add_heading("Zoho CRM Discovery Pack", level=0)
    _docx_set_heading_color(h, "C8102E")

    doc.add_heading(f"Brand: {bundle['brand']}", level=1)
    p = doc.add_paragraph(f"Generated: {bundle['generated_at']}")
    p.runs[0].font.size = Pt(9); p.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # KPI table
    headers = ["Contributors", "Fields", "Mandatory %", "Integration %", "Approvals", "Dashboards"]
    values  = [str(metrics["contributors"]), str(metrics["total_fields"]),
               f"{metrics['mandatory_pct']}%", f"{metrics['integration_pct']}%",
               str(metrics["approval_stages"]), str(metrics["dashboards"])]
    kpi = doc.add_table(rows=2, cols=len(headers)); kpi.style = "Light Grid Accent 1"
    for i, v in enumerate(values):
        cell = kpi.rows[0].cells[i]; cell.text = ""
        run = cell.paragraphs[0].add_run(v)
        run.bold = True; run.font.size = Pt(16); run.font.color.rgb = RGBColor(0xC8, 0x10, 0x2E)
        kpi.rows[1].cells[i].text = headers[i]

    doc.add_paragraph()
    try:
        doc.add_picture(BytesIO(A.chart_completion_donut_png(metrics)), width=Inches(2.4))
        doc.add_picture(BytesIO(A.chart_readiness_radar_png(metrics)),  width=Inches(4.0))
    except Exception:
        pass

    # Contributors
    doc.add_heading("Contributors", level=1)
    if bundle["contributors"]:
        tbl = doc.add_table(rows=1, cols=4); tbl.style = "Light Grid Accent 1"
        for i, h in enumerate(["Name", "Role", "Email", "Last update"]):
            tbl.rows[0].cells[i].text = h
        for c in bundle["contributors"]:
            cells = tbl.add_row().cells
            cells[0].text = c["name"]; cells[1].text = c["role"]
            cells[2].text = c["email"]; cells[3].text = c["submitted_at"][:19].replace("T", " ")
    else:
        doc.add_paragraph("No contributors yet.")

    # Executive summary
    doc.add_page_break()
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        f"This pack captures the Zoho CRM discovery for the {bundle['brand']} practice. "
        f"{metrics['contributors']} contributor(s) proposed {metrics['total_fields']} fields, "
        f"{metrics['mandatory_pct']}% mandatory and {metrics['integration_pct']}% needing integration. "
        f"Data sourced from {metrics['unique_sources']} system(s). "
        f"{metrics['approval_stages']} approval stage(s) and {metrics['dashboards']} dashboard(s) requested. "
        f"Completion {metrics['completion_pct']}%, brand-readiness {A.overall_readiness_pct(metrics):.0f}/100."
    )
    try:
        doc.add_picture(BytesIO(A.chart_fields_per_section_png(bundle)),     width=Inches(6.5))
        doc.add_picture(BytesIO(A.chart_mandatory_pct_png(bundle)),          width=Inches(6.5))
        doc.add_page_break()
        doc.add_heading("Integration & Data Sources", level=1)
        doc.add_picture(BytesIO(A.chart_integration_per_section_png(bundle)),width=Inches(6.5))
        doc.add_picture(BytesIO(A.chart_source_breakdown_png(bundle)),       width=Inches(5.5))
    except Exception:
        pass

    # Conflicts
    if conflicts:
        doc.add_page_break()
        h = doc.add_heading("Conflicts to Resolve", level=1); _docx_set_heading_color(h, "C8102E")
        tbl = doc.add_table(rows=1, cols=4); tbl.style = "Light Grid Accent 1"
        for i, hh in enumerate(["Section", "Type", "Item", "Detail"]):
            tbl.rows[0].cells[i].text = hh
        for c in conflicts:
            row = tbl.add_row().cells
            row[0].text = c["section"]; row[1].text = c["type"]
            row[2].text = c["item"];    row[3].text = c["detail"]

    # Recommendations
    doc.add_page_break()
    doc.add_heading("Insights & Recommendations", level=1)
    doc.add_paragraph("Auto-generated. Color-coded: High = act before handover, Medium = working session, Low = consider, Info = healthy.")
    tbl = doc.add_table(rows=1, cols=3); tbl.style = "Light Grid Accent 1"
    for i, hh in enumerate(["Priority", "Topic", "Recommendation"]):
        tbl.rows[0].cells[i].text = hh
    for r in recs:
        row = tbl.add_row().cells
        row[0].text = r["priority"]; row[1].text = r["topic"]; row[2].text = r["text"]

    # Per-section detail
    for sec_key, sec_title in SECTIONS:
        doc.add_page_break()
        doc.add_heading(sec_title, level=1)
        any_data = False
        for c in bundle["contributors"]:
            payload = bundle["responses_by_contributor"].get(c["id"], {}).get(sec_key)
            if not payload: continue
            any_data = True
            doc.add_heading(f"From {c['name']} ({c['role']})", level=2)
            _render_section_docx(doc, sec_key, payload)
        if not any_data:
            doc.add_paragraph("No input captured for this section.")

    buf = BytesIO(); doc.save(buf)
    return buf.getvalue()


def _render_section_docx(doc: Document, sec_key, payload):
    if sec_key in FIELD_BUILDER_SECTIONS:
        rows = payload.get("fields", []) or []
        if rows:
            tbl = doc.add_table(rows=1, cols=7); tbl.style = "Light Grid Accent 1"
            for i, h in enumerate(["Field", "Type", "Mand.", "Integ.", "Source", "Options", "Conditional Rule"]):
                tbl.rows[0].cells[i].text = h
            for f in rows:
                cells = tbl.add_row().cells
                cells[0].text = _safe(f.get("field"))
                cells[1].text = _safe(f.get("type"))
                cells[2].text = "Yes" if f.get("mandatory") else "No"
                cells[3].text = "Yes" if f.get("integration_needed") else "No"
                cells[4].text = _safe(f.get("data_capture_source"))
                cells[5].text = _safe(f.get("options"))
                cells[6].text = _safe(f.get("conditional_rule"))
        if payload.get("notes"):
            doc.add_paragraph(f"Notes: {_safe(payload['notes'])}")

    elif sec_key == "approvals":
        rows = payload.get("stages", []) or []
        if rows:
            tbl = doc.add_table(rows=1, cols=5); tbl.style = "Light Grid Accent 1"
            for i, h in enumerate(["Stage", "Approver", "Trigger", "SLA (h)", "Can revert?"]):
                tbl.rows[0].cells[i].text = h
            for s in rows:
                cells = tbl.add_row().cells
                cells[0].text = _safe(s.get("stage")); cells[1].text = _safe(s.get("approver"))
                cells[2].text = _safe(s.get("trigger")); cells[3].text = _safe(s.get("sla_hours"))
                cells[4].text = "Yes" if s.get("can_revert") else "No"
        for k in ("escalation", "notes"):
            if payload.get(k):
                doc.add_paragraph(f"{k.title()}: {_safe(payload[k])}")

    elif sec_key == "dashboards":
        rows = payload.get("dashboards", []) or []
        if rows:
            tbl = doc.add_table(rows=1, cols=3); tbl.style = "Light Grid Accent 1"
            for i, h in enumerate(["Dashboard", "Audience", "Frequency"]):
                tbl.rows[0].cells[i].text = h
            for d in rows:
                cells = tbl.add_row().cells
                cells[0].text = _safe(d.get("dashboard"))
                cells[1].text = _safe(d.get("audience"))
                cells[2].text = _safe(d.get("frequency"))
        if payload.get("notes"):
            doc.add_paragraph(f"Notes: {_safe(payload['notes'])}")

    else:
        for k, v in payload.items():
            if v:
                doc.add_paragraph(f"{k.replace('_',' ').title()}: {_safe(v)}")
